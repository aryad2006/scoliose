#!/usr/bin/env python3
"""Convert all CDC_DIAB MD files to DOCX with proper table fitting."""

import os
import glob
import markdown
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re


def md_to_docx(md_path, docx_path, color1=None, color2=None):
    """Convert a markdown file to a formatted DOCX."""
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Set narrow margins for more space
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(12)

    # Colors (customizable per formation)
    BLEU_PETROLE = color1 or RGBColor(0x1A, 0x52, 0x76)
    TURQUOISE = color2 or RGBColor(0x14, 0x8F, 0x77)
    header_hex = "{:02X}{:02X}{:02X}".format(BLEU_PETROLE[0], BLEU_PETROLE[1], BLEU_PETROLE[2])

    lines = content.split("\n")
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Detect table start
        if "|" in stripped and not in_table:
            # Check if this is a table (has | delimiters)
            if stripped.startswith("|") or ("|" in stripped and i + 1 < len(lines) and "---" in lines[i + 1]):
                in_table = True
                table_rows = []
                table_rows.append(stripped)
                i += 1
                continue

        # Continue collecting table rows
        if in_table:
            if stripped.startswith("|") or ("|" in stripped and stripped.count("|") >= 2):
                table_rows.append(stripped)
                i += 1
                continue
            else:
                # End of table, process it
                in_table = False
                process_table(doc, table_rows, header_color=header_hex)
                table_rows = []
                # Don't increment i, process current line normally

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Headers
        if stripped.startswith("######"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.size = Pt(12)
            i += 1
            continue
        elif stripped.startswith("#####"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.size = Pt(13)
            i += 1
            continue
        elif stripped.startswith("####"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.size = Pt(14)
            i += 1
            continue
        elif stripped.startswith("###"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = TURQUOISE
            p.space_before = Pt(14)
            p.space_after = Pt(6)
            i += 1
            continue
        elif stripped.startswith("##"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.size = Pt(17)
            run.font.color.rgb = BLEU_PETROLE
            p.space_before = Pt(18)
            p.space_after = Pt(8)
            i += 1
            continue
        elif stripped.startswith("#"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = BLEU_PETROLE
            p.space_before = Pt(24)
            p.space_after = Pt(10)
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            p = doc.add_paragraph()
            p.add_run("─" * 80).font.size = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Blockquote (MEDIA tags, clinical hooks, key points)
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)

            # Format inline bold/italic
            add_formatted_run(p, text)

            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Bullet lists
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_run(p, text)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            i += 1
            continue

        # Numbered lists
        match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if match:
            text = match.group(2)
            p = doc.add_paragraph(style="List Number")
            add_formatted_run(p, text)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            i += 1
            continue

        # Sub-bullets (indented)
        if stripped.startswith("  - ") or stripped.startswith("  * "):
            text = stripped.strip().lstrip("-*").strip()
            p = doc.add_paragraph(style="List Bullet 2")
            add_formatted_run(p, text)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_run(p, stripped)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        i += 1

    # Process remaining table if any
    if in_table and table_rows:
        process_table(doc, table_rows, header_color=header_hex)

    doc.save(docx_path)


def process_table(doc, table_rows, header_color="1A5276"):
    """Create a table in the document that fits within page width."""
    if len(table_rows) < 2:
        return

    # Parse rows
    parsed_rows = []
    for row_str in table_rows:
        cells = [c.strip() for c in row_str.strip("|").split("|")]
        # Skip separator rows (---)
        if all(re.match(r"^[-:]+$", c.strip()) for c in cells if c.strip()):
            continue
        parsed_rows.append(cells)

    if len(parsed_rows) < 1:
        return

    num_cols = max(len(r) for r in parsed_rows)
    # Normalize all rows to same number of columns
    for r in parsed_rows:
        while len(r) < num_cols:
            r.append("")

    # A4 = 21cm, margins 1.8cm x2 => 17.4cm usable => in twips (1cm = 567 twips)
    avail_twips = int(17.4 * 567)  # ~9866 twips
    col_width_twips = int(avail_twips / num_cols)

    # Create table
    table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Configure table properties
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    # Remove existing tblW and tblLayout to avoid duplicates
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for tag_name in ("tblW", "tblLayout"):
        for el in tbl_pr.findall(f"{{{ns}}}{tag_name}"):
            tbl_pr.remove(el)

    # Table width = exact page width (dxa = twips)
    tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{avail_twips}" w:type="dxa"/>')
    tbl_pr.append(tbl_width)

    # FIXED layout: forces columns to respect specified widths, text wraps
    tbl_layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    tbl_pr.append(tbl_layout)

    for row_idx, row_data in enumerate(parsed_rows):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""

            # Set cell width in twips
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # Remove existing tcW and noWrap
            for tag_name in ("tcW", "noWrap"):
                for el in tcPr.findall(f"{{{ns}}}{tag_name}"):
                    tcPr.remove(el)

            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_width_twips}" w:type="dxa"/>')
            tcPr.append(tcW)

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            clean_text = cell_text.strip()
            add_formatted_run(p, clean_text)

            # Font size for tables
            for run in p.runs:
                run.font.size = Pt(10)

            # Header row: bold + white text on blue background
            if row_idx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(11)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_color}" w:val="clear"/>')
                tcPr.append(shading)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Small spacing after table
    doc.add_paragraph().paragraph_format.space_before = Pt(4)


def add_formatted_run(paragraph, text):
    """Add text to paragraph with basic markdown formatting (bold, italic, code)."""
    if not text:
        return

    # Use finditer to locate all formatted spans, then fill gaps with plain text
    pattern = re.compile(
        r"\*\*\*(.+?)\*\*\*"   # ***bold italic***
        r"|\*\*(.+?)\*\*"      # **bold**
        r"|\*(.+?)\*"          # *italic*
        r"|`(.+?)`"            # `code`
    )

    last_end = 0
    for m in pattern.finditer(text):
        # Plain text before this match
        if m.start() > last_end:
            plain = text[last_end:m.start()]
            if plain:
                run = paragraph.add_run(plain)
                run.font.size = Pt(12)

        if m.group(1):  # ***bold italic***
            run = paragraph.add_run(m.group(1))
            run.bold = True
            run.italic = True
            run.font.size = Pt(12)
        elif m.group(2):  # **bold**
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.font.size = Pt(12)
        elif m.group(3):  # *italic*
            run = paragraph.add_run(m.group(3))
            run.italic = True
            run.font.size = Pt(12)
        elif m.group(4):  # `code`
            run = paragraph.add_run(m.group(4))
            run.font.name = "Consolas"
            run.font.size = Pt(11)

        last_end = m.end()

    # Remaining plain text after last match
    if last_end < len(text):
        remaining = text[last_end:]
        if remaining:
            run = paragraph.add_run(remaining)
            run.font.size = Pt(12)


def convert_batch(src_dir, dst_dir, pattern, color1=None, color2=None):
    """Convert a batch of MD files to DOCX."""
    os.makedirs(dst_dir, exist_ok=True)
    md_files = sorted(glob.glob(os.path.join(src_dir, pattern)))
    count = 0
    for md_file in md_files:
        basename = os.path.splitext(os.path.basename(md_file))[0]
        docx_file = os.path.join(dst_dir, basename + ".docx")
        print(f"Converting {basename}...", end=" ", flush=True)
        try:
            md_to_docx(md_file, docx_file, color1=color1, color2=color2)
            print("OK")
            count += 1
        except Exception as e:
            print(f"ERROR: {e}")
    return count


def main():
    total = 0

    # Diabetologie
    total += convert_batch(
        r"c:\Users\USER\Documents\scoliose\cours-diabetologie",
        r"c:\Users\USER\docs-diabète",
        "CDC_DIAB_*.md"
    )

    # Histoire de l'orthopedie
    total += convert_batch(
        r"c:\Users\USER\Documents\scoliose\cours-histoire-orthopedie",
        r"c:\Users\USER\docs-histoire-orthopedie",
        "CDC_HISTO_*.md",
        color1=RGBColor(0x8B, 0x1A, 0x1A),  # Bordeaux
        color2=RGBColor(0xB8, 0x86, 0x0B),  # Or antique
    )

    print(f"\nDone! {total} files converted")


if __name__ == "__main__":
    main()
