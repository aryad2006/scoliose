#!/usr/bin/env python3
"""
Conversion des Cas Cliniques Diabétologie → DOCX + ODT
Formatage professionnel avec code couleurs par niveau :
  🥉 Bronze  → #CD7F32
  🥈 Argent  → #5D6D7E
  🥇 Or      → #D4AC0D
  💎 Diamant → #76449B
"""

import os, re
from pathlib import Path

# ── DOCX ──────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── ODT ───────────────────────────────────────────────────────────────────────
from odf.opendocument import OpenDocumentText
from odf.style import (Style, TextProperties, ParagraphProperties,
                       TableColumnProperties, TableCellProperties,
                       TableProperties)
from odf.text import H, P, List, ListItem, Span
from odf.table import Table, TableColumn, TableRow, TableCell
from odf import teletype

# ══════════════════════════════════════════════════════════════════════════════
# PALETTE
# ══════════════════════════════════════════════════════════════════════════════
C = {
    'bleu':    RGBColor(0x1A, 0x52, 0x76),
    'teal':    RGBColor(0x14, 0x8F, 0x77),
    'bronze':  RGBColor(0xCD, 0x7F, 0x32),
    'argent':  RGBColor(0x5D, 0x6D, 0x7E),
    'or':      RGBColor(0xD4, 0xAC, 0x0D),
    'diamant': RGBColor(0x76, 0x44, 0x9B),
    'vert':    RGBColor(0x1E, 0x8B, 0x4C),
    'rouge':   RGBColor(0xAB, 0x27, 0x12),
    'orange':  RGBColor(0xD3, 0x58, 0x00),
}
HEX = {k: f'{v[0]:02X}{v[1]:02X}{v[2]:02X}' for k, v in C.items()}
HEX.update({'blanc': 'FFFFFF', 'gris': 'F4F6F7', 'bleu_clair': 'EBF5FB'})

LEVEL_KEYWORDS = {
    'bronze': ('bronze', '🥉'),
    'argent': ('argent', 'silver', '🥈'),
    'or':     ('or ', 'gold', '🥇', ' or\n', ' or$'),
    'diamant':('diamant', 'diamond', '💎'),
}

def detect_level(text: str) -> str:
    t = text.lower()
    if any(k in t for k in ('💎', 'diamant')):  return 'diamant'
    if any(k in t for k in ('🥇',)):             return 'or'
    if '🥇' in text or re.search(r'\bor\b', t):  return 'or'
    if any(k in t for k in ('🥈', 'argent')):    return 'argent'
    if any(k in t for k in ('🥉', 'bronze')):    return 'bronze'
    return ''

def level_color(level: str):
    return C.get(level, C['bleu'])

def level_hex(level: str):
    return HEX.get(level, HEX['bleu'])

# ══════════════════════════════════════════════════════════════════════════════
# DOCX HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def add_run_fmt(para, text: str, base_size=12):
    """Parse **bold**, *italic*, ***both***, `code` inline markup."""
    pat = re.compile(
        r'\*\*\*(.+?)\*\*\*'
        r'|\*\*(.+?)\*\*'
        r'|\*(.+?)\*'
        r'|`(.+?)`'
    )
    last = 0
    for m in pat.finditer(text):
        if m.start() > last:
            r = para.add_run(text[last:m.start()])
            r.font.size = Pt(base_size)
        if m.group(1):
            r = para.add_run(m.group(1)); r.bold = r.italic = True
        elif m.group(2):
            r = para.add_run(m.group(2)); r.bold = True
        elif m.group(3):
            r = para.add_run(m.group(3)); r.italic = True
        elif m.group(4):
            r = para.add_run(m.group(4)); r.font.name = 'Consolas'; r.font.size = Pt(10)
        else:
            r = para.add_run(m.group(0))
        r.font.size = Pt(base_size)
        last = m.end()
    if last < len(text):
        r = para.add_run(text[last:]); r.font.size = Pt(base_size)


def build_docx_table(doc, rows_raw):
    parsed = []
    for row_str in rows_raw:
        cells = [c.strip() for c in row_str.strip('|').split('|')]
        if all(re.match(r'^[-:]+$', c) for c in cells if c):
            continue
        parsed.append(cells)
    if not parsed:
        return
    ncols = max(len(r) for r in parsed)
    for r in parsed:
        while len(r) < ncols: r.append('')

    avail = int(17.4 * 567)
    cw    = avail // ncols
    ns    = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    tbl = doc.add_table(rows=len(parsed), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    tp = tbl._tbl.tblPr
    for tag in ('tblW', 'tblLayout'):
        for el in tp.findall(f'{{{ns}}}{tag}'): tp.remove(el)
    tp.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="{avail}" w:type="dxa"/>'))
    tp.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))

    for ri, row_data in enumerate(parsed):
        row = tbl.rows[ri]
        for ci, txt in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ''
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for tag in ('tcW', 'noWrap'):
                for el in tcPr.findall(f'{{{ns}}}{tag}'): tcPr.remove(el)
            tcPr.append(parse_xml(f'<w:tcW {nsdecls("w")} w:w="{cw}" w:type="dxa"/>'))
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            add_run_fmt(p, txt.strip(), base_size=10)
            if ri == 0:
                for r in p.runs:
                    r.bold = True; r.font.size = Pt(11)
                    r.font.color.rgb = C['blanc'] if hasattr(C, 'blanc') else RGBColor(0xFF,0xFF,0xFF)
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX["bleu"]}" w:val="clear"/>')
                tcPr.append(shd)
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif ri % 2 == 0:
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX["bleu_clair"]}" w:val="clear"/>')
                tcPr.append(shd)
    doc.add_paragraph().paragraph_format.space_before = Pt(4)


def md_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding='utf-8')
    doc  = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(1.8)
        sec.right_margin  = Cm(1.8)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)

    lines       = text.split('\n')
    i           = 0
    in_table    = False
    table_rows  = []
    cur_level   = ''   # bronze / argent / or / diamant

    while i < len(lines):
        line    = lines[i]
        stripped = line.strip()

        # ── table detection ──────────────────────────────────────────────────
        if '|' in stripped and not in_table:
            if stripped.startswith('|') or (i + 1 < len(lines) and '---' in lines[i+1]):
                in_table   = True
                table_rows = [stripped]
                i += 1; continue

        if in_table:
            if stripped.startswith('|') or stripped.count('|') >= 2:
                table_rows.append(stripped)
                i += 1; continue
            else:
                in_table = False
                build_docx_table(doc, table_rows)
                table_rows = []

        if not stripped:
            i += 1; continue

        # ── update level context ─────────────────────────────────────────────
        lv = detect_level(stripped)
        if lv:
            cur_level = lv

        # ── headings ─────────────────────────────────────────────────────────
        if stripped.startswith('######'):
            p = doc.add_paragraph()
            r = p.add_run(stripped.lstrip('#').strip())
            r.bold = True; r.font.size = Pt(11)
            i += 1; continue

        elif stripped.startswith('#####'):
            p = doc.add_paragraph()
            r = p.add_run(stripped.lstrip('#').strip())
            r.bold = True; r.font.size = Pt(12)
            i += 1; continue

        elif stripped.startswith('####'):
            p = doc.add_paragraph()
            r = p.add_run(stripped.lstrip('#').strip())
            r.bold = True; r.font.size = Pt(13)
            if cur_level:
                r.font.color.rgb = level_color(cur_level)
            i += 1; continue

        elif stripped.startswith('###'):
            p = doc.add_paragraph()
            r = p.add_run(stripped.lstrip('#').strip())
            r.bold = True; r.font.size = Pt(14)
            lv2 = detect_level(stripped) or cur_level
            r.font.color.rgb = level_color(lv2) if lv2 else C['teal']
            p.space_before = Pt(14); p.space_after = Pt(6)
            i += 1; continue

        elif stripped.startswith('##'):
            p = doc.add_paragraph()
            r = p.add_run(stripped.lstrip('#').strip())
            r.bold = True; r.font.size = Pt(17)
            lv2 = detect_level(stripped) or cur_level
            r.font.color.rgb = level_color(lv2) if lv2 else C['bleu']
            p.space_before = Pt(18); p.space_after = Pt(8)
            i += 1; continue

        elif stripped.startswith('#'):
            p = doc.add_paragraph()
            r = p.add_run(stripped.lstrip('#').strip())
            r.bold = True; r.font.size = Pt(20)
            r.font.color.rgb = C['bleu']
            p.space_before = Pt(24); p.space_after = Pt(10)
            i += 1; continue

        # ── HR ───────────────────────────────────────────────────────────────
        if stripped == '---':
            p = doc.add_paragraph()
            p.add_run('─' * 80).font.size = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1; continue

        # ── blockquote (MEDIA / clés) ────────────────────────────────────────
        if stripped.startswith('>'):
            txt = stripped.lstrip('>').strip()
            p   = doc.add_paragraph()
            p.paragraph_format.left_indent    = Cm(0.8)
            p.paragraph_format.space_before   = Pt(2)
            p.paragraph_format.space_after    = Pt(2)
            # fond gris clair via shading XML on paragraph (via run color)
            add_run_fmt(p, txt)
            for r in p.runs:
                r.font.color.rgb = C['argent']
                r.italic = True
            i += 1; continue

        # ── bullet list ──────────────────────────────────────────────────────
        if re.match(r'^(\s{2,})?[-*]\s', stripped):
            indent = len(line) - len(line.lstrip())
            style_name = 'List Bullet 2' if indent >= 2 else 'List Bullet'
            txt = re.sub(r'^[-*]\s', '', stripped)
            p   = doc.add_paragraph(style=style_name)
            add_run_fmt(p, txt)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            i += 1; continue

        # ── numbered list ────────────────────────────────────────────────────
        m = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_run_fmt(p, m.group(2))
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            i += 1; continue

        # ── normal paragraph ─────────────────────────────────────────────────
        p = doc.add_paragraph()
        add_run_fmt(p, stripped)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        i += 1

    if in_table and table_rows:
        build_docx_table(doc, table_rows)

    doc.save(str(docx_path))
    print(f'  ✓ DOCX → {docx_path.name}')


# ══════════════════════════════════════════════════════════════════════════════
# ODT CONVERSION
# ══════════════════════════════════════════════════════════════════════════════
LEVEL_ODT_STYLE = {
    'bronze':  'BronzeHeading',
    'argent':  'ArgentHeading',
    'or':      'OrHeading',
    'diamant': 'DiamantHeading',
}

def hex_odt(rgb: RGBColor) -> str:
    return f'#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}'

def make_odt_styles(doc: OpenDocumentText):
    """Register all paragraph/text styles."""
    def para_style(name, font_size, bold=False, italic=False,
                   color=None, margin_top=0, margin_bottom=0,
                   left_margin=0, bg=None):
        s = Style(name=name, family='paragraph')
        pp_attrs = {}
        if margin_top:    pp_attrs['margintop']    = f'{margin_top}pt'
        if margin_bottom: pp_attrs['marginbottom'] = f'{margin_bottom}pt'
        if left_margin:   pp_attrs['marginleft']   = f'{left_margin}cm'
        if bg:            pp_attrs['backgroundcolor'] = bg
        if pp_attrs:
            s.addElement(ParagraphProperties(**pp_attrs))
        tp_attrs = {'fontsize': f'{font_size}pt', 'fontfamily': 'Calibri'}
        if bold:   tp_attrs['fontweight'] = 'bold'
        if italic: tp_attrs['fontstyle']  = 'italic'
        if color:  tp_attrs['color']      = color
        s.addElement(TextProperties(**tp_attrs))
        doc.styles.addElement(s)

    # Base styles
    para_style('Body',         12)
    para_style('Blockquote',   11, italic=True, color=hex_odt(C['argent']), left_margin=0.8)
    para_style('HR',            8)
    para_style('BulletItem',   12)
    para_style('BulletItem2',  12, left_margin=1.0)
    para_style('NumberedItem', 12)

    # H1 — titre principal
    para_style('H1', 20, bold=True, color=hex_odt(C['bleu']),
               margin_top=24, margin_bottom=10)

    # H2 generic
    para_style('H2', 17, bold=True, color=hex_odt(C['bleu']),
               margin_top=18, margin_bottom=8)

    # H2 level-colored
    for lv, col in [('bronze', C['bronze']), ('argent', C['argent']),
                    ('or', C['or']), ('diamant', C['diamant'])]:
        para_style(f'H2_{lv}', 17, bold=True, color=hex_odt(col),
                   margin_top=18, margin_bottom=8)
        para_style(f'H3_{lv}', 14, bold=True, color=hex_odt(col),
                   margin_top=14, margin_bottom=6)

    # H3 generic
    para_style('H3', 14, bold=True, color=hex_odt(C['teal']),
               margin_top=14, margin_bottom=6)
    para_style('H4', 13, bold=True, margin_top=10, margin_bottom=4)
    para_style('H5', 12, bold=True)
    para_style('H6', 11, bold=True)

    # Table styles
    ts = Style(name='MedTable', family='table')
    ts.addElement(TableProperties(width='17cm', align='center'))
    doc.styles.addElement(ts)

    tcs_header = Style(name='TCHeader', family='table-cell')
    tcs_header.addElement(TableCellProperties(backgroundcolor=f'#{HEX["bleu"]}',
                                               padding='0.1cm'))
    doc.styles.addElement(tcs_header)

    tcs_even = Style(name='TCEven', family='table-cell')
    tcs_even.addElement(TableCellProperties(backgroundcolor=f'#{HEX["bleu_clair"]}',
                                             padding='0.1cm'))
    doc.styles.addElement(tcs_even)

    tcs_odd = Style(name='TCOdd', family='table-cell')
    tcs_odd.addElement(TableCellProperties(padding='0.1cm'))
    doc.styles.addElement(tcs_odd)

    # Text span styles
    sp_bold = Style(name='Bold', family='text')
    sp_bold.addElement(TextProperties(fontweight='bold'))
    doc.styles.addElement(sp_bold)

    sp_italic = Style(name='Italic', family='text')
    sp_italic.addElement(TextProperties(fontstyle='italic'))
    doc.styles.addElement(sp_italic)

    sp_bi = Style(name='BoldItalic', family='text')
    sp_bi.addElement(TextProperties(fontweight='bold', fontstyle='italic'))
    doc.styles.addElement(sp_bi)

    sp_code = Style(name='Code', family='text')
    sp_code.addElement(TextProperties(fontfamily='Courier New', fontsize='10pt'))
    doc.styles.addElement(sp_code)

    sp_white = Style(name='WhiteText', family='text')
    sp_white.addElement(TextProperties(color='#ffffff', fontweight='bold',
                                        fontsize='11pt'))
    doc.styles.addElement(sp_white)


def odt_para_with_markup(doc_or_cell, text: str, style_name: str):
    """Create a P element with inline bold/italic/code spans."""
    p = P(stylename=style_name)
    pat = re.compile(
        r'\*\*\*(.+?)\*\*\*'
        r'|\*\*(.+?)\*\*'
        r'|\*(.+?)\*'
        r'|`(.+?)`'
    )
    last = 0
    for m in pat.finditer(text):
        if m.start() > last:
            p.addText(text[last:m.start()])
        if m.group(1):
            sp = Span(stylename='BoldItalic'); sp.addText(m.group(1)); p.addElement(sp)
        elif m.group(2):
            sp = Span(stylename='Bold'); sp.addText(m.group(2)); p.addElement(sp)
        elif m.group(3):
            sp = Span(stylename='Italic'); sp.addText(m.group(3)); p.addElement(sp)
        elif m.group(4):
            sp = Span(stylename='Code'); sp.addText(m.group(4)); p.addElement(sp)
        last = m.end()
    if last < len(text):
        p.addText(text[last:])
    return p


def build_odt_table(odt_doc, rows_raw):
    parsed = []
    for row_str in rows_raw:
        cells = [c.strip() for c in row_str.strip('|').split('|')]
        if all(re.match(r'^[-:]+$', c) for c in cells if c):
            continue
        parsed.append(cells)
    if not parsed:
        return
    ncols = max(len(r) for r in parsed)
    for r in parsed:
        while len(r) < ncols: r.append('')

    col_w = f'{17/ncols:.2f}cm'
    tbl = Table(stylename='MedTable')
    for _ in range(ncols):
        cs = Style(name=f'TC_W{_}', family='table-column')
        cs.addElement(TableColumnProperties(columnwidth=col_w))
        odt_doc.styles.addElement(cs)
        tbl.addElement(TableColumn(stylename=f'TC_W{_}'))

    for ri, row_data in enumerate(parsed):
        tr = TableRow()
        tbl.addElement(tr)
        for ci, txt in enumerate(row_data):
            if ri == 0:
                cell_style = 'TCHeader'
            elif ri % 2 == 0:
                cell_style = 'TCEven'
            else:
                cell_style = 'TCOdd'
            tc = TableCell(stylename=cell_style)
            tr.addElement(tc)
            text_style = 'WhiteText' if ri == 0 else 'Body'
            p = odt_para_with_markup(tc, txt.strip(), text_style)
            tc.addElement(p)

    odt_doc.text.addElement(tbl)
    odt_doc.text.addElement(P())   # spacing after


def md_to_odt(md_path: Path, odt_path: Path):
    text     = md_path.read_text(encoding='utf-8')
    odt_doc  = OpenDocumentText()
    make_odt_styles(odt_doc)

    lines      = text.split('\n')
    i          = 0
    in_table   = False
    table_rows = []
    cur_level  = ''

    while i < len(lines):
        line     = lines[i]
        stripped = line.strip()

        # table
        if '|' in stripped and not in_table:
            if stripped.startswith('|') or (i+1 < len(lines) and '---' in lines[i+1]):
                in_table = True; table_rows = [stripped]; i += 1; continue

        if in_table:
            if stripped.startswith('|') or stripped.count('|') >= 2:
                table_rows.append(stripped); i += 1; continue
            else:
                in_table = False
                build_odt_table(odt_doc, table_rows)
                table_rows = []

        if not stripped:
            i += 1; continue

        lv = detect_level(stripped)
        if lv: cur_level = lv

        # headings
        if stripped.startswith('######'):
            txt = stripped.lstrip('#').strip()
            odt_doc.text.addElement(odt_para_with_markup(odt_doc, txt, 'H6'))
        elif stripped.startswith('#####'):
            txt = stripped.lstrip('#').strip()
            odt_doc.text.addElement(odt_para_with_markup(odt_doc, txt, 'H5'))
        elif stripped.startswith('####'):
            txt = stripped.lstrip('#').strip()
            sn = f'H4'
            odt_doc.text.addElement(odt_para_with_markup(odt_doc, txt, sn))
        elif stripped.startswith('###'):
            txt = stripped.lstrip('#').strip()
            lv2 = detect_level(stripped) or cur_level
            sn  = f'H3_{lv2}' if lv2 else 'H3'
            odt_doc.text.addElement(odt_para_with_markup(odt_doc, txt, sn))
        elif stripped.startswith('##'):
            txt = stripped.lstrip('#').strip()
            lv2 = detect_level(stripped) or cur_level
            sn  = f'H2_{lv2}' if lv2 else 'H2'
            odt_doc.text.addElement(odt_para_with_markup(odt_doc, txt, sn))
        elif stripped.startswith('#'):
            txt = stripped.lstrip('#').strip()
            odt_doc.text.addElement(odt_para_with_markup(odt_doc, txt, 'H1'))
        elif stripped == '---':
            odt_doc.text.addElement(P(stylename='HR'))
            p = P(stylename='HR'); p.addText('─' * 80)
            odt_doc.text.addElement(p)
        elif stripped.startswith('>'):
            txt = stripped.lstrip('>').strip()
            odt_doc.text.addElement(odt_para_with_markup(odt_doc, txt, 'Blockquote'))
        elif re.match(r'^(\s{2,})?[-*]\s', stripped):
            indent = len(line) - len(line.lstrip())
            txt = re.sub(r'^[-*]\s', '', stripped)
            sn  = 'BulletItem2' if indent >= 2 else 'BulletItem'
            p   = odt_para_with_markup(odt_doc, '• ' + txt, sn)
            odt_doc.text.addElement(p)
        elif re.match(r'^(\d+)\.\s+(.*)', stripped):
            m = re.match(r'^(\d+)\.\s+(.*)', stripped)
            p = odt_para_with_markup(odt_doc, f'{m.group(1)}. {m.group(2)}', 'NumberedItem')
            odt_doc.text.addElement(p)
        else:
            odt_doc.text.addElement(odt_para_with_markup(odt_doc, stripped, 'Body'))

        i += 1

    if in_table and table_rows:
        build_odt_table(odt_doc, table_rows)

    odt_doc.save(str(odt_path))
    print(f'  ✓ ODT  → {odt_path.name}')


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════
def main():
    src_dir = Path('/Users/imac/Desktop/scoliose/cours-diabetologie')
    dst_dir = Path('/Users/imac/Desktop/docs-scoliose')
    dst_dir.mkdir(parents=True, exist_ok=True)

    # Uniquement les fichiers CAS_CLINIQUES
    md_files = sorted(src_dir.glob('CAS_CLINIQUES_DIABETOLOGIE_*.md'))
    if not md_files:
        print('Aucun fichier CAS_CLINIQUES_DIABETOLOGIE_*.md trouvé.')
        return

    print(f'\n{"═"*60}')
    print(f'  Conversion vers : {dst_dir}')
    print(f'  Fichiers source : {len(md_files)}')
    print(f'{"═"*60}\n')

    for md in md_files:
        print(f'▶ {md.name}')
        base = md.stem
        try:
            md_to_docx(md, dst_dir / f'{base}.docx')
        except Exception as e:
            print(f'  ✗ DOCX ERREUR: {e}')
        try:
            md_to_odt(md, dst_dir / f'{base}.odt')
        except Exception as e:
            print(f'  ✗ ODT  ERREUR: {e}')
        print()

    files = list(dst_dir.glob('CAS_CLINIQUES_DIABETOLOGIE_*'))
    print(f'{"═"*60}')
    print(f'  ✅  {len(files)} fichiers générés dans {dst_dir}')
    print(f'{"═"*60}\n')


if __name__ == '__main__':
    main()
